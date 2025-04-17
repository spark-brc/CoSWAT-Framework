#!/bin/python3

'''
This module is a collection of functions and
classes which Celray James frequently uses with
Python 3.11+

Author  : Celray James CHAWANDA
Email   : celray.chawanda@outlook.com
Licence : All rights Reserved
Repo    : https://github.com/celray

Updated : 26/12/2022
'''

import os
import re
import sys
import urllib
from collections import defaultdict
import ctypes
import base64
import datetime
import gzip
import math
import string
import multiprocessing
import pickle
import platform
import random
import shutil
import sqlite3
import subprocess
import sys
import time
from io import StringIO
import warnings
import xml.etree.ElementTree as ET
import zipfile
from functools import partial
from glob import glob
from itertools import product
from shutil import copyfile, move, rmtree
from sqlite3.dbapi2 import DatabaseError


def install_package(package_name, version = None):
    if platform.system() == "Windows":
        try: os.system(f"pip3 install {package_name}{'==' + version if not version is None else ''}")
        except: raise

    else:
        install_command = f"pip3 install {package_name}{'==' + version if not version is None else ''} --break-system-packages"
        print(install_command)
        os.system(install_command)
        

try: from pytube import YouTube, Playlist
except: install_package("pytube"); from pytube import YouTube, Playlist

try: import geopandas
except: install_package("geopandas"); import geopandas

try: import matplotlib.pyplot as plt
except: install_package("matplotlib"); import matplotlib.pyplot as plt

try: import numpy
except: install_package("numpy"); import numpy
import hydroeval

try: import pandas
except: install_package("pandas"); import pandas

try: import psycopg2
except: install_package("psycopg2")
    
try: import psycopg2
except: install_package("psycopg2-binary"); import psycopg2

try: import requests
except: install_package("requests"); import requests

try: import rtree
except: install_package("rtree"); import rtree

try: import seaborn
except: install_package("seaborn"); import seaborn

try: import xarray
except: install_package("xarray"); import xarray

try: import xlsxwriter; 
except: install_package("xlsxwriter"); import xlsxwriter
from xlsxwriter.utility import xl_rowcol_to_cell

try: from PIL import Image
except: install_package("pillow"); from PIL import Image

try: from pySmartDL import SmartDL
except: install_package("pySmartDL"); from pySmartDL import SmartDL

try: from tqdm import tqdm
except: install_package("tqdm"); from tqdm import tqdm

try: from wand import image
except:
    if platform.system() == "Windows":
        response = requests.get('https://imagemagick.org/archive/binaries/ImageMagick-7.1.1-4-Q16-HDRI-x64-dll.exe', stream=True)
        dst_dir = f"{os.path.dirname(__file__)}/__pycache__/"
        if not os.path.isdir(dst_dir): os.makedirs(dst_dir)

        with open(f"{dst_dir}/ImageMagick-7.1.1-4-Q16-HDRI-x64-dll.exe", "wb") as handle:
            for data in tqdm(response.iter_content()):
                handle.write(data)
        os.system(f"{dst_dir}/ImageMagick-7.1.1-4-Q16-HDRI-x64-dll.exe")
    
    install_package("wand"); from wand import image

try: from osgeo import gdal, gdalconst, ogr
except:
    if platform.system() == "Windows":
        print('downloading and installing gdal')
        import sys, genericpath

        urls_by_v = {
            '3.8'   : 'https://celray.chawanda.com/assets/downloads/GDAL-3.4.3-cp38-cp38-win_amd64.whl',
            '3.9'   : 'https://celray.chawanda.com/assets/downloads/GDAL-3.4.3-cp39-cp39-win_amd64.whl',
            '3.10'  : 'https://celray.chawanda.com/assets/downloads/GDAL-3.4.3-cp310-cp310-win_amd64.whl',
            '3.11'  : 'https://celray.chawanda.com/assets/downloads/GDAL-3.4.3-cp311-cp311-win_amd64.whl',
        }
        
        dst_dir = f"{os.path.dirname(__file__)}/__pycache__/"
        if not os.path.isdir(dst_dir): os.makedirs(dst_dir)
        obj = SmartDL(urls_by_v[f"{sys.version_info.major}.{sys.version_info.minor}"], dst_dir)
        try:
            obj.start()
            path = obj.get_dest()
        except: raise
        os.system(f'pip install "{path}"')

    else:
        os.system("sudo pacman -S gdal-bin")
        os.system("sudo pacman -Sy libgdal-dev")

        install_package("gdal")

try: from osgeo import gdal, gdalconst, ogr
except: import gdal, gdalconst, ogr

try: import wget
except: install_package("wget"); import wget

try: import rasterio
except: install_package("rasterio"); import rasterio

try: import openpyxl
except: install_package("openpyxl"); import openpyxl

try: import docx
except: install_package("python-docx"); import docx
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

try: import pyodbc
except: pyodbc = install_package("pyodbc"); import pyodbc

try: import sqlalchemy
except: sqlalchemy = install_package("sqlalchemy"); import sqlalchemy
from sqlalchemy import MetaData, Table, create_engine, func, select

try: import pydub
except: install_package("pydub"); import pydub
from pydub import AudioSegment

from shapely.geometry import Point, Polygon
from shapely import wkt




docstring = \
'''
Purpose

Author  : Celray James
Email   : celray.chawanda@outlook.com
Date    : 00/00/2021
Licence : MIT (2021 or later)
'''


def points_to_geodataframe(point_pairs_list, columns = ['latitude', 'longitude'], auth = "EPSG", code = '4326', out_shape = '', format = 'gpkg', v = False, get_geometry_only = False):
    df = pandas.DataFrame(point_pairs_list, columns = columns)
    geometry = [Point(xy) for xy in zip(df['latitude'], df['longitude'])]

    if get_geometry_only:
        return geometry[0]

    gdf = geopandas.GeoDataFrame(point_pairs_list, columns = columns, geometry=geometry)
    drivers = {'gpkg': 'GPKG', 'shp': 'ESRI Shapefile'}

    gdf = gdf.set_crs(f'{auth}:{code}')
    
    if out_shape != '':
        if v: print(f'creating shapefile {out_shape}')
        gdf.to_file(out_shape, driver=drivers[format])
    
    return gdf


def isYearInFileRange(fileName, yearToCheck):
    # match all sequences of 4 digits (potential years)
    possibleYears = re.findall(r'(?<!\d)(\d{4})(?!\d)', fileName)
    
    if len(possibleYears) < 2:
        raise ValueError("Could not find two years in the file name.")

    # convert to integers and sort
    possibleYears = sorted([int(year) for year in possibleYears])
    startYear, endYear = possibleYears[0], possibleYears[-1]

    return startYear <= int(yearToCheck) <= endYear

def get_relative_path(base_path, target_path):
    return os.path.relpath(target_path, base_path)

def get_usgs_nutrients(file_path, monthly=True):

    if not exists(file_path):
        df = pandas.DataFrame(columns=['date', 'loadAMLE'])
        df = df.astype({'date': 'datetime64[ns]', 'loadAMLE': 'float64'})
        return df

    df = pandas.read_csv(file_path, delim_whitespace=True, names=['date', 'flow', 'loadAMLE', 'loadMLE', "loadLAD"], skiprows=1)
    df['date'] = pandas.to_datetime(df['date'], format='%Y%m')
    
    if monthly:
        df = df.groupby(pandas.Grouper(key='date', freq='M')).sum()
        df.reset_index(inplace=True)
        df.rename(columns={'index': 'date'}, inplace=True)

    return df

def alert(message, message_title = "info", attachment = None, priority = None, tags = [], server = "http://ntfy.chawanda.com", print_it = True, topic = "pythonAlerts", v = False):
    '''
    This sends an alert to a given server in case you want to be notified of something
    '''
    print(message) if print_it else None; header_data = {}
    if not message_title is None: header_data["Title"] = message_title
    if not priority is None: header_data["Priority"] = priority
    if not len(tags) == 0: header_data["Tags"] = ",".join(tags)

    try:
        if v: print(f"sending alert to {server}/{topic}")
        if not attachment is None:
            header_data["Filename"] = file_name(attachment)
            requests.put( f"{server}/{topic}", data=open(attachment, 'rb'), headers=header_data )
        try: requests.post(f"{server}/{topic}",data=message, headers=header_data )
        except: pass
    except: pass

def get_usgs_timeseries(file_path, monthly=False, convert=True, var_column_name='flow'):
    df = pandas.read_csv(file_path, delim_whitespace=True, names=['date', 'Additional_1', var_column_name, 'Additional_2'])
    df['date'] = pandas.to_datetime(df['date'], format='%Y%m%d')
    
    if convert:
        df[var_column_name] = df[var_column_name] * 0.0283168466 

    if monthly:
        df = df.groupby(pandas.Grouper(key='date', freq='M')).mean()
        df.reset_index(inplace=True)
        df.rename(columns={'index': 'date'}, inplace=True)

    return df[['date', var_column_name]]



class dual_progress_bar:
    def __init__(self, main_goal, secondary_goal, ncols = 100):
        self.main_goal = main_goal
        self.secondary_goal = secondary_goal
        self.ncols = ncols

        self.main_pbar = tqdm(total=self.main_goal, ncols=self.ncols, bar_format='{l_bar}{bar}| {percentage:3.0f}%')
        self.secondary_pbar = tqdm(total=self.secondary_goal, ncols=self.ncols, bar_format='{l_bar}{bar}| {percentage:3.0f}%')

    def update(self, secondary_counter, main_counter):
        self.main_pbar.update(main_counter - self.main_pbar.n)
        self.secondary_pbar.update(secondary_counter - self.secondary_pbar.n)
        self.secondary_pbar.refresh()
        sys.stdout.write('\033[1A')  # cursor up one line
        sys.stdout.write('\033[2K')  # erase to end of line
        self.main_pbar.refresh()

    def set_secondary_goal(self, new_secondary_goal):
        self.secondary_goal = new_secondary_goal
        self.secondary_pbar.total = self.secondary_goal

    def close(self):
        self.main_pbar.close()
        self.secondary_pbar.close()
    
    def main_message(self, message:str)->None:
        print("\n" + message)
        return None


def decode_64(encoded:str) -> str:
    return base64.b64decode(encoded).decode('utf-8')

def wait(duration:int = 1):
    time.sleep(duration)

def ignore_warnings(ignore:bool = True, v:bool = False):
    if ignore:
        warnings.filterwarnings("ignore")
        if v: print("warnings ignored")
    else:
        warnings.filterwarnings("default")
        if v: print("warnings not ignored")

def get_raster_value_for_coords(lat, lon, rasterio_array):
    row, col = rasterio_array.index(lon,lat)
    val = rasterio_array.read(1)[row,col]
    return(val)

def rasterise_shape(shape_file:str, column_name:str, destination_tif:str, template_tif:str, no_data:int = -999) -> None:
    data = gdal.Open(template_tif, gdalconst.GA_ReadOnly)
    geo_transform = data.GetGeoTransform()
    x_min = geo_transform[0]
    y_max = geo_transform[3]
    x_max = x_min + geo_transform[1] * data.RasterXSize
    y_min = y_max + geo_transform[5] * data.RasterYSize
    x_res = data.RasterXSize
    y_res = data.RasterYSize
    mb_v = ogr.Open(shape_file)
    mb_l = mb_v.GetLayer()
    pixel_width = geo_transform[1]
    target_ds = gdal.GetDriverByName('GTiff').Create(destination_tif, x_res, y_res, 1, gdal.GDT_Int32)
    target_ds.SetGeoTransform(data.GetGeoTransform())
    target_ds.SetProjection(data.GetProjection())
    # target_ds.SetGeoTransform((x_min, pixel_width, 0, y_min, 0, pixel_width))
    band = target_ds.GetRasterBand(1)
    band.SetNoDataValue(no_data)
    band.FlushCache()
    gdal.RasterizeLayer(target_ds, [1], mb_l, options=[f"ATTRIBUTE={column_name}"])

    target_ds = None

def extract_timeseries_from_netcdf(fn, variable, lon_coord, lat_coord, method = 'nearest', v = False):
    if v: print(f'getting a {lon_coord},{lat_coord} timeseries for {variable}')
    ds = xarray.open_dataset(fn)
    dsloc = ds.sel(lon=lon_coord, lat=lat_coord, method=method)
    return dsloc[variable].to_dataframe()

def rand_apha_num(length = 8):
    return ''.join(random.choices(string.ascii_letters + string.digits, k=length))

def make_plot(data_pd, x_col:str, y1_cols:list, y1_axis_label, save_path:str, title = None, y2_axis_label = None, y2_cols:list = None,
                x_axis_label:str=None, y1_labels:list = None, y2_labels:list = None, x_limits:tuple = None, y1_limits:tuple = None, y2_limits:tuple = None,
                show_plot = False, legend = False, secondary_type = 'bar', chart_size:tuple = (14, 5)):
    '''
    x_limits : example (0, 10)
    y_labels : example ['ET', 'PCP'] # no need to be same as y_cols but y_cols are used if this is not given

    secondary_type: 'bar', 'line'
    '''

    import seaborn
    from matplotlib import pyplot as plt

    plt.close()
    plt.margins(x=None, y=None, tight=True)
    plt.tight_layout(pad = 0)

    ax1 = seaborn.set_style(style=None, rc=None )

    fig = plt.figure(figsize = chart_size, dpi=300)

    # make axis
    ax1  = fig.add_subplot(1, 1, 1)

    min_ax1_ini =  1000000000000
    max_ax1_ini = -1000000000000
    for index in range(0, len(y1_cols)):
        ax1.plot(data_pd[x_col].tolist(), data_pd[y1_cols[index]].tolist(), label = y1_labels[index] if not y1_labels is None else y1_cols[index])

        if min(data_pd[y1_cols[index]].tolist()) < min_ax1_ini: min_ax1_ini = min(data_pd[y1_cols[index]].tolist())
        if max(data_pd[y1_cols[index]].tolist()) > max_ax1_ini: max_ax1_ini = max(data_pd[y1_cols[index]].tolist())

    min_ax1 = max(0, min_ax1_ini - (max_ax1_ini - min_ax1_ini)/5 )
    max_ax1 = max_ax1_ini + (max_ax1_ini - min_ax1_ini)/1.2

    min_ax1 = round(min_ax1/10)*10 if min_ax1 < 80 else (round(min_ax1/100)*100 if min_ax1 < 800 else round(min_ax1/1000)*1000)

    if min_ax1 > min_ax1_ini:
        difference_ax1 = round(1 + (min_ax1 - min_ax1_ini)/10) * 10 if (min_ax1 - min_ax1_ini) < 80 else (round(1 + (min_ax1 - min_ax1_ini)/100)*100 if (min_ax1 - min_ax1_ini) < 800 else round(1 + (min_ax1 - min_ax1_ini)/1000)*1000)
        min_ax1 = min_ax1 - difference_ax1

    difference_ax1 = round(1 + (max_ax1 - min_ax1)//10) * 10 if (max_ax1 - min_ax1) < 80 else (round(1 + (max_ax1 - min_ax1)//100)*100 if (max_ax1 - min_ax1) < 800 else round(1 + (max_ax1 - min_ax1)/1000)*1000)

    max_ax1 = min_ax1 + difference_ax1

    # set axes labels
    ax1.set_ylabel(y1_axis_label)
    if not title is None: ax1.set_title(title, fontdict={'fontsize': 18})
    if not x_axis_label is None: ax1.set_xlabel(x_axis_label)

    if y1_limits is None:
        ax1.set_ylim([min_ax1, max_ax1])
        ax1.grid(axis='y', which='major', color = 'lightgrey')
        ax1.set_yticks([min_ax1 + (i * (max_ax1 - min_ax1) / 5) for i in range(1, 6)])
    else: ax1.set_ylim([y1_limits[0], y1_limits[1]])

    if not x_limits is None: ax1.set_xlim([x_limits[0], x_limits[1]])


    if not y2_cols is None:
        ax2 = ax1.twinx()

        min_ax2_ini =  1000000000000
        max_ax2_ini = -1000000000000

        if secondary_type == 'bar':
            for index in range(0, len(y2_cols)):
                ax2.bar(data_pd[x_col].tolist(), data_pd[y2_cols[index]].tolist(), 0.6, label = y2_labels[index] if not y2_labels is None else y2_cols[index], color = 'black')
                
                if min(data_pd[y2_cols[index]].tolist()) < min_ax2_ini: min_ax2_ini = min(data_pd[y2_cols[index]].tolist())
                if max(data_pd[y2_cols[index]].tolist()) > max_ax2_ini: max_ax2_ini = max(data_pd[y2_cols[index]].tolist())
        
        elif secondary_type == 'line':
            for index in range(0, len(y2_cols) + 1):
                ax2.plot(data_pd[x_col].tolist(), data_pd[y2_cols[index]].tolist(), label = y2_labels[index] if not y2_labels is None else y2_cols[index])
                
                if min(data_pd[y2_cols[index]].tolist()) < min_ax2_ini: min_ax2_ini = min(data_pd[y2_cols[index]].tolist())
                if max(data_pd[y2_cols[index]].tolist()) > max_ax2_ini: max_ax2_ini = max(data_pd[y2_cols[index]].tolist())

        else: raise ValueError('\'secondary_type\' can only be either \'bar\' or \'line\'')

        # label axes
        if not y2_axis_label is None: ax2.set_ylabel(y2_axis_label)

        min_ax2 = max(0, min_ax2_ini - (max_ax2_ini - min_ax2_ini)/5)
        max_ax2 = max_ax2_ini + (max_ax2_ini - min_ax2_ini)/0.9

        min_ax2 = round(min_ax2/10)*10 if min_ax2 < 80 else (round(min_ax2/100)*100 if min_ax2 < 800 else round(min_ax2/1000)*1000)

            
        if min_ax2 > min_ax2_ini:
            difference_ax2 = round(1 + (min_ax2 - min_ax2_ini)/10) * 10 if (min_ax2 - min_ax2_ini) < 80 else (round(1 + (min_ax2 - min_ax2_ini)/100)*100 if (min_ax2 - min_ax2_ini) < 800 else round(1 + (min_ax2 - min_ax2_ini)/1000)*1000)
            min_ax2 = min_ax2 - difference_ax2

        difference_ax2 = round(1 + (max_ax2 - min_ax2)//10) * 10 if (max_ax2 - min_ax2) < 80 else (round(1 + (max_ax2 - min_ax2)//100)*100 if (max_ax2 - min_ax2) < 800 else round(1 + (max_ax2 - min_ax2)/1000)*1000)

        max_ax2 = min_ax2 + difference_ax2




        # invert axis (need to set limits too)
        if y2_limits is None:
            ax2.set_yticks([i * (max_ax2 - min_ax2) / 5 for i in range(1, 6)])
            ax2.set_ylim([min_ax2, max_ax2][::-1])
        else:
            ax2.set_ylim([y2_limits[0], y2_limits[1]][::-1])

    handles,labels = [],[]
    for ax in fig.axes:
        for h,l in zip(*ax.get_legend_handles_labels()):
            handles.append(h)
            labels.append(l)

    if legend: plt.legend(handles, labels, loc="center right",)

    plt.savefig(save_path)
    if show_plot: plt.show()

    return plt

def goto_dir(obj_):
    '''
    obj_ should be __file__
    '''
    me = os.path.realpath(obj_)
    os.chdir(os.path.dirname(me))

def smart_copy(source_dir, destination_dir, move = False):
    destination_dir = destination_dir if not destination_dir.endswith("/") else destination_dir[:-1]
    source_dir = source_dir if not source_dir.endswith("/") else source_dir[:-1]
    splitter = os.path.basename(source_dir)

    all_files_names = list_all_files(source_dir)
    class file_paths:
        def __init__(self, src, dst):
            self.src = src
            self.dst = dst

    sort_size = []
    files_dict = {}

    print("\n\tsorting files")
    for fn in all_files_names:
        size = os.path.getsize(fn)
        while size in sort_size:
            size += random.random()
        sort_size.append(size)
        files_dict[size] = file_paths(src = fn, dst = "{0}/{1}{2}".format(destination_dir, splitter, fn.split(splitter)[-1]))

    sort_size.sort(reverse=True)

    count = 0
    end = len(sort_size)
    for sorted_file in sort_size:
        count += 1
        show_progress(count = count, end_val=end, string_after= "{2} {0} - {1} MB".format(
            os.path.basename(files_dict[sorted_file].dst).lower(),
            round(sorted_file/1000000, 3),
            "copying" if not move else "moving"
            ))

        if not os.path.isfile(files_dict[sorted_file].src):
            print("\nfile not found")
            return

        copy_file(files_dict[sorted_file].src, files_dict[sorted_file].dst, delete_source = move)

    print("")

def heic_to_jpg(heic_path:str, dst_path:str, v:bool = True) -> str:
    if v:
        print(f"processing {file_name(heic_path)}")
    
    create_path(dst_path)
    img=image.Image(filename=heic_path)
    img.format='jpg'
    img.save(filename=dst_path)
    img.close()
    
    return dst_path

def print_dict(dictionary, columns = 4, width_k = 10, width_val = 12, sep = ":") -> None:
    counter = columns

    keys = [k for k in dictionary]
    keys.sort()

    print_str = ""
    for key in keys:
        print_str += f"{str(key).rjust(width_k)} {sep} {str(dictionary[key]).ljust(width_val)}     "
        counter -= 1

        if counter == 0:
            counter = columns
            print(print_str)
            print_str = ""
        
    print(print_str)
    return print_str


def transparent_image(in_file, out_file, threshold = 251, v = True):
    img = Image.open(in_file)
    img = img.convert("RGBA")

    pixdata = img.load()

    width, height = img.size
    for y in range(height):
        for x in range(width):
            if pixdata[x, y][0] > threshold:
                if pixdata[x, y][1]  > threshold:
                    if pixdata[x, y][2]  > threshold:
                        pixdata[x, y] = (255, 255, 255, 0)

    img.save(out_file, "PNG")
    if v:
        print(f"\t\t - saved transparent image to {out_file}")

def get_file_size(file_path):
    return float(os.path.getsize(file_path))/1012

def cd(path):
    if os.path.isdir(path):
        os.chdir(path)
    else:
        sys.exit("\t! the path {0} does not exist".format(path))

def slope_intercept(x1,y1,x2,y2):
    a = (y2 - y1) / (x2 - x1)
    b = y1 - a * x1     
    return a,b

def distance(coords_a, coords_b):
    return math.sqrt(((coords_b[0] - coords_a[0]) ** 2) + (coords_b[1] - coords_a[1]) ** 2)

def merge_documents(out_fn, list_of_in_fn, from_empty = False, v = False):
    if len(list_of_in_fn) < 2:
        print('\t! you need atleast two documents')
        return

    from docx import Document
    from docxcompose.composer import Composer
    
    create_path(out_fn)

    if from_empty:
        # initialise empty document
        document = Document(); document.save(out_fn)
    else:
        if not list_of_in_fn[0] == out_fn:
            copy_file(list_of_in_fn[0], out_fn)
        else:
            if v: print(f"\t ! appending to the existing out file")

        list_of_in_fn = list_of_in_fn[1:]

    master = Document(out_fn)
    composer = Composer(master)

    for doc_fn in list_of_in_fn:
        composer.append(Document(doc_fn))

    master.save(out_fn)  

def resize_image(in_file:str, out_file:str, size:int = None, ratio = None, width = None, height = "auto"):
    fsize = get_file_size(in_file)

    if not size is None:
        ratio = 1

        class point:
            def __init__(self, X, Y) -> None:
                self.X = X
                self.Y = Y

        history = []
        # print(abs(fsize-size))
        while fsize > size:
            # if len(history) > 1:
            #     a, b = slope_intercept(history[-2].X, history[-2].Y, history[-1].X, history[-1].Y)
            #     ratio = b/a
            # else:
            #     ratio -= 0.2
            # if ratio < 0.05: ratio = random.randint(1, 9) * random.randint(1, 9) * 0.01
            image = Image.open(in_file)

            old_width, old_height = image.size
            newsize = (int(old_width * ratio), int(old_height * ratio))

            image = image.resize(newsize)
            create_path(out_file)
            image.save(out_file)

            fsize = get_file_size(out_file)
            ratio -= 0.01

            sys.stdout.write('\r\t> testing ratio: ' + str(ratio))
            sys.stdout.flush()

            history.append(point(ratio, fsize-size))

        print(f'\n\t> reduced by {ratio}\n')
        
    else:
        image = Image.open(in_file)

        old_width, old_height = image.size

        if ratio is None:
            if width is None:
                print("You should supply either width or ratio")
                return

            if height == "auto":
                newsize = (width, int((old_height/old_width) * width))
            else:
                newsize = (width, height)
            
        else:
            newsize = (int(old_width * ratio), int(old_height * ratio))

        image = image.resize(newsize)
        create_path(out_file)
        image.save(out_file)
        return True
    
def download(url, dest_dir, replace = False, v = True):

    create_path(dest_dir)
    if v: print(f"downloading {url}...")
    dest_dir = f"{dest_dir}/" if not dest_dir.endswith("/") else dest_dir
    
    obj = SmartDL(url, dest_dir)
    
    if not replace:
        path = obj.get_dest()

        if exists(path):
            return path
        if exists(f"{dest_dir}{url.split('/')[-1]}".replace('%20', ' ')):
            return path

    try:
        obj.start()
        path = obj.get_dest()
        return path
    except:
        print(f"\t! could not download {url}")
        write_to(f"{dest_dir}failed.txt", url, mode="a")
        return None


import os
import requests
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor
import math

def download_chunk(url, start, end, path):
    headers = {'Range': f'bytes={start}-{end}'}
    response = requests.get(url, headers=headers, stream=True)
    with open(path, 'wb') as f:
        for chunk in response.iter_content(chunk_size=8192):
            if chunk:
                f.write(chunk)

def download_file(url, save_path, exists_action='resume', num_connections=5, v=True):
    if v:
        print(f"\ndownloading {url}")
    fname = file_name(url, extension=True)
    save_dir = os.path.dirname(save_path)
    save_fname = "{0}/{1}".format(save_dir, fname)

    if not os.path.isdir(save_dir):
        os.makedirs(save_dir)

    # Handle existing file
    if os.path.exists(save_fname):
        if exists_action == 'skip':
            if v:
                print(f"File exists, skipping: {save_fname}")
            return
        elif exists_action == 'overwrite':
            os.remove(save_fname)
        # 'resume' is handled below

    # Get file size
    response = requests.head(url)
    file_size = int(response.headers.get('content-length', 0))

    # Resume download if file exists and exists_action is 'resume'
    initial_pos = 0
    if exists_action == 'resume' and os.path.exists(save_fname):
        initial_pos = os.path.getsize(save_fname)
        if initial_pos >= file_size:
            if v:
                print(f"File already completed: {save_fname}")
            return

    # Calculate chunk sizes
    chunk_size = math.ceil((file_size - initial_pos) / num_connections)
    chunks = []
    for i in range(num_connections):
        start = initial_pos + (i * chunk_size)
        end = min(start + chunk_size - 1, file_size - 1)
        chunks.append((start, end))

    # Download chunks in parallel
    temp_files = [f"{save_fname}.part{i}" for i in range(num_connections)]
    with ThreadPoolExecutor(max_workers=num_connections) as executor:
        futures = []
        for i, (start, end) in enumerate(chunks):
            futures.append(
                executor.submit(download_chunk, url, start, end, temp_files[i])
            )
        
        # Wait for all downloads to complete with progress bar
        with tqdm(total=file_size-initial_pos, initial=initial_pos, unit='B', 
                 unit_scale=True, desc=fname) as pbar:
            completed = initial_pos
            while completed < file_size:
                current = sum(os.path.getsize(f) for f in temp_files if os.path.exists(f))
                pbar.update(current - completed)
                completed = current

    # Merge chunks
    with open(save_fname, 'ab' if initial_pos > 0 else 'wb') as outfile:
        for temp_file in temp_files:
            if os.path.exists(temp_file):
                with open(temp_file, 'rb') as infile:
                    outfile.write(infile.read())
                os.remove(temp_file)


def apply_parallel(function_name, number_of_processes, *args):
    "has issues"
    with multiprocessing.Pool(processes=number_of_processes) as pool:
        results = pool.starmap(function_name, product(args[0], args[1]))

    return results


def xml_children_attributes(xml_file_name, x_path):
    root = ET.parse(xml_file_name).getroot()
    result = {}
    for element in root.findall(x_path):
        for child in element:
            result[child.tag] = child.text

    return result


def python_variable(option, filename, variable=None):
    '''
    option: save, load or open

    '''
    if (option == "save") and (variable is None):
        print("\t! please specify a variable")

    if option == "save":
        create_path(filename)
        with open(filename, 'wb') as f:
            pickle.dump(variable, f)

    if (option == "load") or (option == "open"):
        with open(filename, "rb") as f:
            variable = pickle.load(f)

    return variable

def get_kge(data: pandas.DataFrame, observed_column: str, simulated_column: str):
    '''
    This function calculates the Kling-Gupta Efficiency (KGE) from columns in a pandas
    dataframe. The dataframe should have two columns, one for the simulated data and
    one for the observed data.
    '''
    df_clean = data.dropna(subset=[simulated_column, observed_column])
    series1 = df_clean[observed_column]
    series2 = df_clean[simulated_column]
    r = numpy.corrcoef(series1, series2)[0, 1]
    alpha = numpy.std(series1) / numpy.std(series2)
    beta = numpy.mean(series1) / numpy.mean(series2)
    kge = 1 - numpy.sqrt((r - 1)**2 + (alpha - 1)**2 + (beta - 1)**2)

    return kge

def get_nse(data: pandas.DataFrame, observed_column: str, simulated_column: str):
    """
    This function calculates the NSE from columns in a pandas dataframe. The dataframe
    should have two columns, one for the simulated data and one for the observed data.
    The columns should be named as the arguments of this function. The dataframe may not
    have a date but it is recommended to have a date column.
    """
    df_clean = data.dropna(subset=[simulated_column, observed_column])
    try:
        nse = 1 - sum((df_clean[observed_column] - df_clean[simulated_column]) ** 2) / \
            sum((df_clean[observed_column] - numpy.mean(df_clean[observed_column])) ** 2)
    except ZeroDivisionError:
        nse = numpy.nan

    return nse

def get_re(data: pandas.DataFrame, observed_column: str, simulated_column: str):
    """
    This function calculates the relative error from columns in a pandas dataframe. The dataframe
    should have two columns, one for the simulated data and one for the observed data.
    The columns should be named as the arguments of this function. The dataframe may not
    have a date but it is recommended to have a date column.
    """
    df_clean = data.dropna(subset=[simulated_column, observed_column])
    try:
        relative_error = (df_clean[observed_column] - df_clean[simulated_column]).abs() / df_clean[observed_column]
    except ZeroDivisionError:
        relative_error = numpy.nan

    return relative_error.mean()



def get_pbias(data: pandas.DataFrame, observed_column: str, simulated_column: str):
    """
    This function calculates the PBIAS from columns in a pandas dataframe. The dataframe
    should have two columns, one for the simulated data and one for the observed data.
    The columns should be named as the arguments of this function. The dataframe may not
    have a date but it is recommended to have a date column.
    """
    df_clean = data.dropna(subset=[simulated_column, observed_column])
    try:
        pbias = sum((df_clean[observed_column] - df_clean[simulated_column])) / \
            sum(df_clean[observed_column]) * 100
    except ZeroDivisionError:
        pbias = None
    return pbias


class postgres_connection:
    def __init__(self, host = "localhost", database = "db", user = "script", password = "script"):
        self.connection     = None
        self.host           = host
        self.database       = database
        self.user           = user
        self.password       = password
        self.cursor         = None
    
    def connect(self):
        self.connection = psycopg2.connect(
            host = self.host,
            database = self.database,
            user = self.user,
            password = self.password
        )
        self.cursor = self.connection.cursor()

    def execute_query(self, query):
        if self.cursor is None:
            report("the database is closed")
            return None
        
        self.cursor.execute(query)


    def update_value(self):
        pass

    def create_table(self, table_name, column_names, column_types):
        sql_command = f"CREATE TABLE {table_name}(\n"
        for i in range(0, len(column_names)):
            sql_command += f"   {column_names[i].ljust(16)} {column_types[i]}{'' if i == len(column_names) - 1 else ','}\n"
        sql_command += ")"

        self.execute_query(sql_command)
        return True


    def rename_table(self):
        pass
    def table_exists(self):
        pass
    def delete_rows(self):
        pass
    def delete_table(self):
        pass
    def read_table_columns(self):
        pass
    def insert_field(self):
        pass
    def insert_row(self):
        pass
    def insert_rows(self):
        pass
    def dump_csv(self):
        pass

    def commit(self):
        self.connection.commit()

    def close_connection(self, commit = False):
        if commit:
            self.commit()

        self.cursor.close()
        self.connection.close()




class sqlite_connection:
    def __init__(self, sqlite_database, connect = False):
        self.db_name = sqlite_database
        self.connection = None
        self.cursor = None

        if connect:
            self.connect()

    def connect(self, v=True):
        self.connection = sqlite3.connect(self.db_name)
        self.cursor = self.connection.cursor()
        if v:
            report("\t-> connection to " + self.db_name + " established...")

    def update_value(self, table_name, col_name, new_value, col_where1, val_1, v=False):
        """
        does not work yet!
        """
        if not new_value is None:
            new_value = str(new_value)
            self.cursor.execute("UPDATE " + table_name + " SET " + col_name +
                                " = '" + new_value + "' WHERE " + col_where1 + " = " + val_1 + ";")
        if new_value is None:
            self.cursor.execute("UPDATE " + table_name + " SET " + col_name +
                                " = ? " + " WHERE " + col_where1 + " = ?", (new_value, val_1))
        # self.cursor.execute(sql_str)
        if v:
            report("\t -> updated {1} value in {0}".format(
                self.db_name.split("/")[-1].split("\\")[-1], table_name))

    def create_table(self, table_name, initial_field_name, data_type):
        '''
        can be text, real, etc
        '''
        try:
            self.cursor.execute('''CREATE TABLE ''' + table_name +
                                '(' + initial_field_name + ' ' + data_type + ')')
            report("\t-> created table " + table_name + " in " + self.db_name)
        except:
            report("\t! table exists")

    def rename_table(self, old_table_name, new_table_name, v=False):
        """
        this function gives a new name to an existing table and saves the changes
        """
        self.cursor.execute("ALTER TABLE " + old_table_name +
                            " RENAME TO " + new_table_name)
        if v:
            report("\t-> renamed " + old_table_name + " to " + new_table_name)
        self.commit_changes()

    def table_exists(self, table_name):
        self.cursor.execute("SELECT count(name) FROM sqlite_master WHERE type='table' AND name='{table_name}'".format(
            table_name=table_name))
        if self.cursor.fetchone()[0] == 1:
            return True
        else:
            return False

    def delete_rows(self, table_to_clean, col_where=None, col_where_value=None, v=False):
        """

        """

        if (col_where is None) and (col_where_value is None):
            self.connection.execute("DELETE FROM " + table_to_clean)

        elif (not col_where is None) and (not col_where_value is None):
            self.connection.execute(
                "DELETE FROM " + table_to_clean + " WHERE " + col_where + " = " + col_where_value + ";")

        else:
            raise ("\t! not all arguments were provided for selective row deletion")

        if v:
            report("\t-> removed all rows from " + table_to_clean)

    def delete_table(self, table_name):
        """
        this function deletes the specified table
        """
        self.cursor.execute('''DROP TABLE ''' + table_name)
        report("\t-> deleted table " + table_name + " from " + self.db_name)

    def undo_changes(self):
        """
        This function reverts the database to status before last commit
        """
        report("\t-> undoing changes to " + self.db_name + " then saving")
        self.connection.rollback()
        self.commit_changes()

    def read_table_dict(self, table_name, key_column = 'id'):
        # Execute a SQL query to fetch all rows from your table
        self.cursor = self.connection.execute(f"SELECT * FROM {table_name}")

        # Fetch all rows as dictionaries
        rows = [dict(zip([column[0] for column in self.cursor.description], row)) for row in self.cursor.fetchall()]

        # Convert the list of dictionaries to a dictionary of dictionaries, 
        # using the 'id' field as the key
        data = {row[key_column]: row for row in rows}

        return data

    def get_columns_with_types(self, table_name):
        c = self.cursor

        # Prepare and execute a PRAGMA table_info statement
        c.execute(f'PRAGMA table_info({table_name})')

        # Fetch all rows and extract the column names and types
        columns_with_types = {row[1]: row[2] for row in c.fetchall()}

        return columns_with_types

    def insert_dict_partial(self, table_name, data_dict):
        c = self.cursor

        # Get the column names from the table
        c.execute(f"PRAGMA table_info({table_name})")
        columns = [row[1] for row in c.fetchall()]

        # Filter the dictionary keys to match the column names
        filtered_data = {k: v for k, v in data_dict.items() if k in columns}

        # Prepare an INSERT INTO statement
        fields = ', '.join(filtered_data.keys())
        placeholders = ', '.join('?' for _ in filtered_data)
        values = list(filtered_data.values())
        sql = f'INSERT INTO {table_name} ({fields}) VALUES ({placeholders})'

        # Execute the statement
        c.execute(sql, values)

        # Commit the changes
        self.commit_changes()



    def create_table_from_dict(self, table_name, columns_with_types):

        # Prepare a CREATE TABLE statement
        fields = ', '.join(f'{column} {data_type}' for column, data_type in columns_with_types.items())
        sql = f'CREATE TABLE IF NOT EXISTS {table_name} ({fields})'

        # Execute the statement
        self.connection.execute(sql)
        self.commit_changes()


    def insert_dict(self, table_name, data):
        
        # Prepare an INSERT INTO statement for each dictionary
        for id, row in data.items():
            fields = ', '.join(row.keys())
            placeholders = ', '.join('?' for _ in row)
            values = list(row.values())
            sql = f'INSERT INTO {table_name} ({fields}) VALUES ({placeholders})'

            # Execute the statement
            self.cursor.execute(sql, values)

        # Commit the changes
        self.connection.commit()



    def read_table_columns(self, table_name, column_list="all"):
        """
        this function takes a list to be a string separated by commmas and
        a table and puts the columns in the table into a variable

        "all" to select all columns
        """
        if column_list == "all":
            self.cursor = self.connection.execute(
                "SELECT * from " + table_name)
        else:
            self.cursor = self.connection.execute(
                "SELECT " + ",".join(column_list) + " from " + table_name)

        list_of_tuples = []
        for row in self.cursor:
            list_of_tuples.append(row)
        self.cursor = self.connection.cursor()
        report("\t-> read selected table columns from " + table_name)
        return list_of_tuples

    def insert_field(self, table_name, field_name, data_type, to_new_line=False, messages=True):
        """
        This will insert a new field into your sqlite database

        table_name: an existing table
        field_name: the field you want to add
        data_type : text, integer, float or real
        """
        self.cursor.execute("alter table " + table_name +
                            " add column " + field_name + " " + data_type)
        if messages:
            if to_new_line:
                report(
                    "\t-> inserted into table {0} field {1}".format(table_name, field_name))
            else:
                sys.stdout.write(
                    "\r\t-> inserted into table {0} field {1}            ".format(table_name, field_name))
                sys.stdout.flush()

    def insert_row(self, table_name, ordered_content_list = [], dictionary_obj = {}, messages=False):
        """
        ordered_list such as ['ha','he','hi']
        list should have data as strings
        """
        if len(ordered_content_list) > 0:
            self.cursor.execute("INSERT INTO " + table_name + " VALUES(" + "'" + "','".join(ordered_content_list) + "'" + ')')

        if len(dictionary_obj) > 0:
            question_marks = ','.join(list('?'*len(dictionary_obj)))
            keys = ','.join(dictionary_obj.keys())
            values = tuple(dictionary_obj.values())
            self.cursor.execute('INSERT INTO '+table_name+' ('+keys+') VALUES ('+question_marks+')', values)

        if messages:
            report("\t-> inserted row into " + table_name)

    def insert_rows(self, table_name, list_of_tuples, messages=False):
        """
        list_of_tuples such as [('ha','he','hi')'
                                ('ha','he','hi')]
        not limited to string data
        """
        self.cursor.executemany('INSERT INTO ' + table_name + ' VALUES (?{qmarks})'.format(
            qmarks=",?" * (len(list_of_tuples[0]) - 1)), list_of_tuples)
        if messages:
            report("\t-> inserted rows into " + table_name)

    def dump_csv(self, table_name, file_name, index=False, v=False):
        '''
        save table to csv
        '''
        tmp_conn = sqlite3.connect(self.db_name)
        df = pandas.read_sql_query(
            "SELECT * FROM {tn}".format(tn=table_name), tmp_conn)
        if index:
            df.to_csv(file_name)
        else:
            df.to_csv(file_name, index=False)

        if v:
            report(
                "\t-> dumped table {0} to {1}".format(table_name, file_name))

    def commit_changes(self, v=False):
        '''
        save changes to the database.
        '''
        self.connection.commit()
        number_of_changes = self.connection.total_changes
        if v:
            report(
                "\t-> saved {0} changes to ".format(number_of_changes) + self.db_name)

    def close_connection(self, commit=True):
        '''
        disconnects from the database
        '''
        if commit:
            self.commit_changes()
        self.connection.close()
        report("\t-> closed connection to " + self.db_name)


class mssql_connection:
    def __init__(self, server = '172.20.255.235', username = 'cjames', password = 'G0dBL3ssTh3USA', driver = '{ODBC Driver 18 for SQL Server}', trust_server_ssl = True) -> None:
        self.server     = server
        self.username   = username
        self.password   = password
        self.driver     = driver

        self.trust_server_ssl     = trust_server_ssl
        
        self.connection = None
        self.cursor     = None
        self.db_name    = "TMP_CJames"
        
        self.databases  = []

    def connect(self):
        # Connect to the SQL Server instance
        connection_string   = f'DRIVER={self.driver};SERVER={self.server};UID={self.username};PWD={self.password};TrustServerCertificate={"yes" if self.trust_server_ssl else "no"}'
        try:
            self.connection = pyodbc.connect(connection_string)
            self.cursor = self.connection.cursor()
            print(f"> connection to {self.server} established...")
        except pyodbc.Error as e:
            print("! error occurred while connecting to the SQL Server instance:")
            print(e)


    def list_databases(self) -> list:
        query = "SELECT name FROM sys.databases"
        
        if self.connection is None:
            print("! there is no connection to the MSSQL server instance")
            sys.exit(1)
            
        # Execute the query and fetch the results
        try:
            self.cursor = self.connection.cursor()
            self.cursor.execute(query)
            self.databases = [row[0] for row in self.cursor.fetchall()]
            print("\n> List of available databases:")
            for db in self.databases:
                print(f"\t- {db}")
        except pyodbc.Error as e:
            print("! error occurred while fetching the list of databases:")
            print(e)
        
        return self.databases
    
    
    def list_tables(self, db_name = None) -> list:
        
        if not db_name is None:
            self.connect_db(db_name)
            
        query = """
        SELECT TABLE_NAME
        FROM INFORMATION_SCHEMA.TABLES
        WHERE TABLE_TYPE = 'BASE TABLE'
        """

        try:
            self.cursor = self.connection.cursor()
            self.cursor.execute(query)
            tables = [row[0] for row in self.cursor.fetchall()]
            print("> list of tables in the active database:")
            for table in tables:
                print(f"\t- {table}")
        except pyodbc.Error as e:
            print("Error occurred while fetching the list of tables:")
            print(e)
        
        return tables
    
        
    def read_table(self, table_name:str, db_name:str = None, columns:list = None, geom_col:str = None, v = True):
        if db_name is not None:
            self.connect_db(db_name)

        if columns is not None and geom_col is not None:
            columns.append(f"{geom_col}.STAsText() as {geom_col}_wkt")
            query = f"SELECT {','.join(columns)} FROM {table_name}"
        elif columns is not None:
            query = f"SELECT {','.join(columns)} FROM {table_name}"
        elif geom_col is not None:
            query = f"SELECT *, {geom_col}.STAsText() as {geom_col}_wkt FROM {table_name}"
        else:
            query = f"SELECT * FROM {table_name}"

        # Load as a regular DataFrame
        if v: print(f"> reading table: {table_name} from {self.db_name}")
        df = pandas.read_sql(query, self.connection)

        # Convert WKT column to a GeoPandas geometry column if needed
        if geom_col is not None:
            df[geom_col] = df[geom_col+"_wkt"].apply(wkt.loads)
            df = geopandas.GeoDataFrame(df, geometry=geom_col)
            
        return df
        
    # Function to change the active database
    def connect_db(self, db_name = None, v = True):
        if not self.connection:
            self.connect()
        try:
            self.cursor     = self.connection.cursor()
            self.cursor.execute(f"USE {db_name if not db_name is None else self.db_name}")
            self.db_name    = db_name if not db_name is None else self.db_name
            self.cursor.commit()

            if v: print(f"> changed active database to: {db_name if not db_name is None else self.db_name}")
        except pyodbc.Error as e:
            print("! error occurred while changing the active database:")
            print(e)

    def dataframe_to_sql(self, df, table_name, if_exists='fail', geom_col='geometry', v = True):
        """
        Write records stored in a DataFrame to a SQL database.
        """
        print(f"> saving data to table: {table_name}...")

        # Create SQLAlchemy engine
        params = urllib.parse.quote_plus(f'DRIVER={self.driver};SERVER={self.server};DATABASE={self.db_name};UID={self.username};PWD={self.password};TrustServerCertificate=yes')
        engine = create_engine(f"mssql+pyodbc:///?odbc_connect={params}")

        # Check if dataframe is a GeoDataFrame and has a geometry column
        if isinstance(df, geopandas.GeoDataFrame) and geom_col in df.columns:
            # Create a new column for WKT format
            df[geom_col+'_wkt'] = df[geom_col].apply(lambda x: x.wkt)
            
            # Drop the original geometry column
            df = df.drop(columns=[geom_col])
            
            # Write DataFrame to SQL table
            df.to_sql(table_name, engine, if_exists=if_exists, index=False)

            # Create a new connection and cursor
            conn = engine.raw_connection()
            cursor = conn.cursor()

            # Convert the WKT column back to a geometry column in SQL Server
            cursor.execute(f"ALTER TABLE [{table_name}] ADD [{geom_col}] geometry")
            cursor.execute(f"UPDATE [{table_name}] SET [{geom_col}] = geometry::STGeomFromText([{geom_col}_wkt], 4326)")
            
            # Drop the WKT column
            cursor.execute(f"ALTER TABLE [{table_name}] DROP COLUMN [{geom_col}_wkt]")
            
            conn.commit()

            # Close the connection and cursor
            cursor.close()
            conn.close()

        else:
            # If dataframe is not a GeoDataFrame or doesn't have a geometry column, write it to SQL as usual
            df.to_sql(table_name, engine, if_exists=if_exists, index=False)

        if v:
            print(f"> saved data to table: {table_name}...")



    def modify_sql_table(self, df, table_name):
        """
        Replace an existing SQL table with a new one based on a DataFrame.

        Parameters:
        df : DataFrame
        table_name : string
            Name of SQL table
        """
        self.dataframe_to_sql(df, table_name, if_exists='replace')
    
    def drop_table(self, table_name, v = True):
        """
        Drops a table from the database.

        Parameters:
        table_name : string
            Name of SQL table
        """
        
        if not self.connection:
            self.connect()

        # Drop the table
        self.cursor.execute(f"DROP TABLE IF EXISTS {table_name}")

        # Commit the transaction
        self.connection.commit()

        if v: print(f"> deleted table {table_name}")

    def delete_table(self, table_name, v = True):
        """
        Drops a table from the database | an alias for drop_table method.

        Parameters:
        table_name : string
            Name of SQL table
        """
        
        if not self.connection:
            self.connect()

        # Drop the table
        self.cursor.execute(f"DROP TABLE IF EXISTS {table_name}")

        # Commit the transaction
        self.connection.commit()

        if v: print(f"> deleted table {table_name}")

    def close(self, v = True):
        if self.connection:
            self.connection.close()
            self.connection = None
            self.cursor = None
        if v: print("> connection closed...")

    def disconnect(self, v = True):
        self.close(v = v)

    def close_connection(self, v = True):
        self.close(v = v)

def report(string, printing=False):
    if printing:
        print(f"\t> {string}")
    else:
        sys.stdout.write("\r" + string)
        sys.stdout.flush()


class word_document:
    def __init__(self, path) -> None:
        self.document = Document()
        self.path     = path if path.endswith('docx') else f'{path}.docx'
    
    def add_heading(self, heading, level = 2):
        self.document.add_heading(heading, level)
    
    def add_paragraph(self, text = "", alignment = 'justify'):
        '''
        allignments:
                - justify,
                - left,
                - center,
                - right,
                - justify-low,
                - justify-med
        '''
        
        p = self.document.add_paragraph(text)
        
        if alignment == 'left':p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        if alignment == 'center':p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if alignment == 'right':p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        if alignment == 'justify':p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        if alignment == 'justify-low':p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_LOW
        if alignment == 'justify-med':p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY_MED


    def add_list_item(self, text = "", numbers = False):
        p = self.document.add_paragraph(text)
        if numbers:
            p.style = 'List Number'
        else:
            p.style = 'List Bullet'
    
    def add_text(self, text, bold = False, italic = False):
        if bold:
            self.document.paragraphs[-1].add_run(text).bold = True
        elif italic:
            self.document.paragraphs[-1].add_run(text).italic = True
        elif bold and italic:
            self.document.paragraphs[-1].add_run(text).bold = True
            self.document.paragraphs[-1].runs[-1].italic = True
        else:
            self.document.paragraphs[-1].add_run(text)


    def add_image(self, path_to_image, width_ = 16):
        self.document.add_picture(path_to_image, width=Cm(width_))

    def add_page_break(self):
        self.document.add_page_break()

    def save(self):
        create_path(self.path)
        self.document.save(self.path)

    def set_margins(self, margin = 1.75):
        sections = self.document.sections
        for section in sections:
            section.top_margin = Cm(margin)
            section.bottom_margin = Cm(margin)
            section.left_margin = Cm(margin)
            section.right_margin = Cm(margin)


class excel:
    def __init__(self, path):
        self.path = path
        self.sheet_names = {}
        self.book = None
        self.chart_names = []
        self.date_format = None

    def create(self):
        create_path(os.path.dirname(self.path))
        self.book = xlsxwriter.Workbook(self.path)

    def add_sheet(self, sheet_name):
        if self.book is None:
            self.create()
        self.sheet_names[sheet_name] = self.book.add_worksheet(sheet_name)

    def set_date_format(self, format_string='dd/mm/yyyy'):
        self.date_format = self.book.add_format({'num_format': format_string})

    def write_date(self, sheet_name, row, column, datetime_obj):
        if self.date_format is None:
            self.set_date_format()
            
        self.sheet_names[sheet_name].write_datetime(
            row, column, datetime_obj, self.date_format)

    def write(self, sheet_name, row, column, value):
        self.sheet_names[sheet_name].write(row, column, value)

    def set_column_width(self, sheet_name, column_names, width=12):
        '''
        column_names: list
        '''
        if isinstance(column_names, str):
            self.sheet_names[sheet_name].set_column(
                "{col}:{col}".format(col=column_names), width)
        else:
            for column in column_names:
                self.sheet_names[sheet_name].set_column(
                    "{col}:{col}".format(col=column), width)

    def add_figure(self, sheet_name, x_src_sheet_name, x_start, x_end, y_src_sheet_name, y_start, y_end, position_cell="E2", chart_type='subtype',
                subtype='straight', title='-', size = [1, 1], width = 720, height = 576, marker_type = 'automatic', x_axis_name = "", y_axis_name = "",
                gridlines_visible = False, xmin = 0, ymin = 0, xmax = None, ymax = None):
        '''
        x_start example : "E3"
        marker_type     : automatic, none, square, diamond, triangle, x, star, short_dash, long_dash, circle, plus
        '''
        chart = self.book.add_chart({'type': 'scatter', })
        chart.set_size({'width': width, 'height': height})
        # axis options
        chart.set_x_axis({
            'name': x_axis_name,
            'min': xmin, 'max': xmax,
            'major_gridlines': {
                    'visible': gridlines_visible,
                },
        })

        chart.set_y_axis({
            'name': y_axis_name,
            'min': ymin, 'max': ymax,
            'major_gridlines': {
                    'visible': gridlines_visible,
                },
        })


        self.sheet_names[sheet_name].insert_chart(
            position_cell, chart, {'x_scale': size[1], 'y_scale': size[0]})
        chart.add_series({
            'categories': '={sht}!{strt_x}:{end_x}'.format(sht=x_src_sheet_name, strt_x=x_start, end_x=x_end),
            'values': '={sht}!{strt_y}:{end_y}'.format(sht=y_src_sheet_name, strt_y=y_start, end_y=y_end),
            'name': title,
            'marker': {'type': marker_type},
            # 'trendline': {'type': 'linear'},
        })
        chart.set_legend({'position': 'bottom'})

    def write_column(self, sheet_name, target_cell, content_list):
        '''target_cell eg = A1'''
        self.sheet_names[sheet_name].write_column(target_cell, content_list)

    def get_platform():
        """Returns: Windows or Linux"""
        return platform.system()


    def start(self):
        if platform.system() == "Windows":
            os.startfile(os.path.abspath(self.path))

    def to_alpha_numeric(self, row, column):
        return xl_rowcol_to_cell(row, column)

    def save(self):
        continue_ = True
        while continue_:
            try:
                self.book.close()
                continue_ = False
            except:
                print("\t! Error writing the Excel file, make sure it is closed")
                answer = input("\t> retry? (y/n): ")
                continue_ = True if answer == "y" else False


def print_list(list_object):
    print("\t> list content")
    for item in list_object:
        print("\t  - {0}".format(item))


def disp(string_):
    print(f"\t- {string_}")

def create_path(path_name, v = False):
    path_name = os.path.dirname(path_name)
    if path_name == '':
        path_name = './'
    if not os.path.isdir(path_name):
        os.makedirs(path_name)
        if v:
            print(f"\t> created path: {path_name}")
    
    return path_name

def flow_duration_curve(df, col_name, pct_col_name = "Exceedance Probability (%)"):
    # Select the data for the specified type
    flow_data = df[col_name]

    # Sort the flow data in decreasing order
    sorted_data = flow_data.sort_values(ascending=False)
    
    # Calculate the rank of each data point
    ranks = numpy.arange(1,len(sorted_data)+1)
    
    # Calculate the exceedance probability
    exceedance_prob = (ranks/(len(sorted_data)+1))*100
    
    # Create a DataFrame for plotting
    df_plot = pandas.DataFrame({
        pct_col_name: exceedance_prob,
        col_name: sorted_data
    }).reset_index(drop=True)
    
    return df_plot


def delete_path(path, v=True):
    try:
        if v:
            rmtree(path)
            print(f"\t! {path} has been deleted")
    except:
        print("\t! could not delete the diractory at {path}".format(path=path))


def is_file(path_):
    if os.path.isfile(path_):
        return True
    else:
        return False


def resample_ts_df(df, column_name, t_step="M", resample_type="mean", only_numeric = True):
    '''
    resample_type   : mean, sum
    t_step          : M - > month
                      Y - > year
                      W - > week
                      3T- > 3 minutes
                      3S- > 3 seconds
    '''
    df[column_name] = pandas.to_datetime(df[column_name])
    if resample_type == "mean":
        result_df = df.resample(t_step, on = column_name).mean(numeric_only = only_numeric)
    elif resample_type == "sum":
        result_df = df.resample(t_step, on = column_name).sum(numeric_only = only_numeric)
    else:
        print("\t ! please select a resample type: available > mean, sum")
        sys.exit()
    return result_df


def raster_statistics(tif_file, v =True):
    ds = gdal.Open(tif_file)
    minimum, maximum, mean, std_dev = ds.GetRasterBand(1).GetStatistics(0, 1)

    class gdal_stats:
        def __init__(self, mn, mx, mean, std_dev):
            self.minimum = mn
            self.maximum = mx
            self.mean = mean
            self.stdev = std_dev

        def __repr__(self):
            return 'min: {0}, max: {1}, mean: {2}, sdev: {3}'.format(self.minimum, self.maximum,  self.mean, self.stdev)
    if v:
        report(f"processing {file_name(tif_file, extension=True)}")

    all_stats = gdal_stats(minimum, maximum, mean, std_dev)
    return all_stats


def empty_line():
    print("")

def hide_folder(path_to_folder):
    ctypes.windll.kernel32.SetFileAttributesW(path_to_folder, 2)

def time_stamp():
    return datetime.datetime.now()

def copy_dir(src, dst):
    if not os.path.exists(dst):
        os.makedirs(dst)
    shutil.copytree(src, dst, dirs_exist_ok=True)

def copy_folder(src_dir, dst_dir, delete_source = False, v = False, progress = True, replace = True, exception_list = None):
    all_files = list_all_files(src_dir)

    current = 0; finish = len(all_files)
    for fn in all_files:
        current += 1
        
        skip_file = False
        if not exception_list is None:
            for exception_item in exception_list:
                if exception_item in fn:
                    skip_file = True

        new_fn = f"{dst_dir}/{get_relative_path(src_dir, fn)}"

        if progress and v:
            if skip_file:
                show_progress(current, finish, scroll_text=f"{'skipping'} {fn}")
            else:
                show_progress(current, finish, scroll_text=f"{'copying' if not delete_source else 'moving'} {fn}")
                copy_file(fn, new_fn, replace=replace, v = False, delete_source=delete_source)
        elif progress:
            if not skip_file:
                show_progress(current, finish)
                copy_file(fn, new_fn, replace=replace, v = False, delete_source=delete_source)

    if delete_source:
        delete_path(src_dir, v=False)


def delete_file(file_name, v = True):
    if os.path.isfile(file_name):
        try:
            os.remove(file_name)
            if v: print("\t> deleted file {fn}".format(fn=file_name))
        except:
            print("\t! could not delete file {fn}".format(fn=file_name))
    else:
        if v: print("\t! the file, {fn}, does not exist".format(fn=file_name))


def run_swat_plus(txtinout_dir, final_dir = os.path.abspath(os.getcwd()), executable_path = '/home/cjames/Nextcloud/modules/.bins/SWATPlus64_linux' if platform.system() == "Linux" else "C:/SWAT/executables/Rev_60_5_6_64rel.exe", v = True, direct = False):
    os.chdir(txtinout_dir)

    # get the directory name without the whole path
    base_dir = os.path.basename(os.path.normpath(txtinout_dir))
    
    if direct:
        os.system(f"{executable_path}")
    else:
        if not v:
            # Run the SWAT+ but ignore output and errors
            subprocess.run([executable_path], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        else:
            
            yrs_line = read_from('time.sim')[2].strip().split()

            yr_from = int(yrs_line[1])
            yr_to = int(yrs_line[3])

            delta = datetime.datetime(yr_to, 12, 31) - datetime.datetime(yr_from, 1, 1)

            CREATE_NO_WINDOW = 0x08000000

            if platform.system() == "Windows":
                process = subprocess.Popen(executable_path, stdout=subprocess.PIPE, creationflags=CREATE_NO_WINDOW )
            else:
                process = subprocess.Popen(executable_path, stdout=subprocess.PIPE)

            counter = 0

            current = 0
            number_of_days = delta.days + 1

            day_cycle = []
            previous_time = None

            while True:
                line = process.stdout.readline()
                # if "mkd" in str(line):
                #     continue
                line_parts = str(line).strip().split()
                if not "Simulation" in line_parts: pass
                elif 'Simulation' in line_parts:
                    ref_index = str(line).strip().split().index("Simulation")
                    year = line_parts[ref_index + 3]
                    month = line_parts[ref_index + 1]
                    day = line_parts[ref_index + 2]


                    month = f"0{month}" if int(month) < 10 else month
                    day = f"0{day}" if int(day) < 10 else day
                    
                    current += 1
                    
                    if not previous_time is None:
                        day_cycle.append(datetime.datetime.now() - previous_time)

                    if len(day_cycle) > 40:
                        if len(day_cycle) > (7 * 365.25):
                            del day_cycle[0]

                        av_cycle_time = sum(day_cycle, datetime.timedelta()) / len(day_cycle)
                        eta = av_cycle_time * (number_of_days - current)

                        eta_str = f"  ETA - {format_timedelta(eta)}:"

                    else:
                        eta_str = ''

                    show_progress(current, number_of_days, bar_length=20, string_before=f"      ", string_after= f' >> [{base_dir}] current date: {day}/{month}/{year} - final-date: 31/12/{yr_to} {eta_str}')

                    previous_time = datetime.datetime.now()
                elif "ntdll.dll" in line_parts:
                    print("\n! there was an error running SWAT+\n")
                if counter < 10:
                    counter += 1
                    continue

                if len(line_parts) < 2: break

            show_progress(1, 1, string_before=f"      ", string_after= f'                                                                                             ')
            print("\n    > SWAT+ simulation complete\n")
        
    os.chdir(final_dir)


def open_file_in_code(file_path):
    '''
    This function opens a file in the code editor
    '''
    if file_path == "":
        file_path = os.path.join(os.getcwd(), __file__)
    os.system(f"code {file_path}")



def format_timedelta(delta: datetime.timedelta) -> str:
    """Formats a timedelta duration to [N days] %H:%M:%S format"""
    seconds = int(delta.total_seconds())

    secs_in_a_day = 86400
    secs_in_a_hour = 3600
    secs_in_a_min = 60

    days, seconds = divmod(seconds, secs_in_a_day)
    hours, seconds = divmod(seconds, secs_in_a_hour)
    minutes, seconds = divmod(seconds, secs_in_a_min)

    time_fmt = f"{hours:02d}:{minutes:02d}:{seconds:02d}"

    if days > 0:
        suffix = "s" if days > 1 else ""
        return f"{days} day{suffix} {time_fmt}"
    else:
        return f"{time_fmt}"


def show_progress(progress, end, dt=None, string_before="", string_after="", bar_length=100, precision = 1, d_count = None, scroll_text = None):
    '''
    dt: timedelta from datetime
    '''

    if platform.system() != "Windows":
        if (os.get_terminal_size(0)[0] - (len(string_after) + 21)) < bar_length:
            bar_length = os.get_terminal_size(0)[0] - (len(string_after) + 21)

    if bar_length < 5: bar_length = 5

    if not scroll_text is None:

        if platform.system() != "Windows":
            sys.stdout.write('\r' + str(scroll_text).ljust(os.get_terminal_size(0)[0]))
        else:
            sys.stdout.write('\r' + str(scroll_text).ljust(max(151, 0)))

        sys.stdout.flush(); print()

    percent = float(progress) / end
    hashes = "" * int(round(percent * bar_length))
    spaces = '' * (bar_length - len(hashes))
    eta = 0
    if dt is not None:

        if len(dt) > 5:

            cycle_time = sum(dt, datetime.timedelta()) / (len(dt) if d_count is None else d_count)
            cycles_to_go = end - progress
            eta = format_timedelta((cycle_time * cycles_to_go))

        else:
            dt = None

    sys.stdout.write("\r{str_b}{bar} {sp}{pct}% {str_after}          ".format(
        str_b=string_before, sp = "  " if percent < 100 else "",
        bar=hashes + spaces,
        pct='{:06.2f}'.format(percent * 100),
        str_after=string_after if dt is None else string_after + " - " + "eta: " + eta))
    sys.stdout.flush()

def equation_of_line(x1, y1, x2, y2):
    if x1 == x2:
        raise ValueError("Infinite slope, the line is vertical")
    
    m = (y2 - y1) / (x2 - x1)
    c = y1 - m * x1
    return m, c

def confirm():
    print("\n\t>> set line reached, terminating program")
    answer = input(
        "\t>> press ENTER to continue, 'N' and ENTER to terminate: ")
    if answer == "N":
        sys.exit()


def open_tif_as_array(tif_file, big_tif = True, band = 1):
    if big_tif:
        dataset = gdal.Open(tif_file, gdal.GA_ReadOnly)
        band = dataset.GetRasterBand(band)
        return band.ReadAsArray()
    else:
        im = Image.open(tif_file)
        imarray = numpy.array(im)
        return imarray


def save_array_as_image(np_array, out_file, v = False):
    if not os.path.dirname(out_file) == "":
        create_path(os.path.dirname(out_file))
    im = Image.fromarray(np_array)
    if v:
        print(f'\t> writing array to {out_file}')
    im.save(out_file)

def create_icon(source_image_path, destination_path):
    # create_path(os.path.dirname(destination_path))
    img = Image.open(source_image_path)
    icon_sizes = [(16, 16), (32, 32), (48, 48), (64, 64)]
    img.save(destination_path if destination_path.endswith(".ico")
             else f"{destination_path}.ico", sizes=icon_sizes)

def unzip_file(f_name, destination_of_contents):
    show("extracting {0} to {1}".format(f_name, destination_of_contents))
    
    create_path(f'{destination_of_contents}/{file_name(f_name, extension=True)[:-3]}')
    if f_name.lower().endswith('.gz'):
        with gzip.open(f_name, 'rb') as f_in:
            with open(f'{destination_of_contents}/{file_name(f_name, extension=True)[:-3]}', 'wb') as f_out:
                shutil.copyfileobj(f_in, f_out)
    else:
        try:
            with zipfile.ZipFile(f_name, "r") as zip_ref:
                zip_ref.extractall(destination_of_contents)
        except:
            print(f"{f_name} is probably a bad zip file")


def single_spaces(string_):
    for i in range(0, 20):
        string_ = string_.replace("  ", " ")
    return string_

def view(file_name, ask = False):
    import platform
    
    if platform.system() == "Windows":
        if not ask:
            os.startfile(file_name)
        else:
            answer = input(f"\t Open {file_name} ?\n\t Y/N/E: ")
            if answer == "Y":
                os.startfile(file_name)
            if answer == "N":pass
            if answer == "E":
                sys.exit()

def quit():
    sys.exit("! mannually exited...")

def insert_newlines(input_str, max_chars=70):
    words = input_str.split(' ')
    lines = []
    current_line = ""

    for word in words:
        # If adding the next word to the current line would exceed the max_chars limit
        if len(current_line) + len(word) > max_chars:
            # Append current line to lines and start a new one
            lines.append(current_line.strip())
            current_line = word
        else:
            # Add the word to the current line
            current_line += " " + word

    # Append any remaining words
    lines.append(current_line.strip())

    return '\n'.join(lines)



def write_to(filename, text_to_write, v=False, mode = "overwrite"):
    '''
    a function to write to file
    modes: overwrite/o; append/a
    '''
    try:
        if not os.path.isdir(os.path.dirname(filename)):
            os.makedirs(os.path.dirname(filename))
            if v:
                print("! the directory {0} has been created".format(
                    os.path.dirname(filename)))
    except:
        pass

    if (mode == "overwrite") or (mode == "o"):
        g = open(filename, 'w', encoding="utf-8")
    elif (mode == "append") or (mode == "a"):
        g = open(filename, 'a', encoding="utf-8")
    try:
        g.write(text_to_write)
        if v:
            print('\n\t> file saved to ' + filename)
    except PermissionError:
        print("\t> error writing to {0}, make sure the file is not open in another program".format(
            filename))
        response = input("\t> continue with the error? (Y/N): ")
        if response == "N" or response == "n":
            sys.exit()
    g.close
    return filename


def exists(path_):
    if os.path.isdir(path_):
        return True
    if os.path.isfile(path_):
        return True
    return False

def gdal_datatypes():
    return {
            'Byte': gdal.GDT_Byte,
            'Int16': gdal.GDT_Int16,
            'Int32': gdal.GDT_Int32,
            'UInt16': gdal.GDT_UInt16,
            'UInt32': gdal.GDT_UInt32,
            'CInt16': gdal.GDT_CInt16,
            'CInt32': gdal.GDT_CInt32,
            'Float32': gdal.GDT_Float32,
            'Float64': gdal.GDT_Float64,
            'CFloat32': gdal.GDT_CFloat32,
            'CFloat64': gdal.GDT_CFloat64,
        }

def resample_raster(original_file:str, destination_file:str, resolution:float, authority = "ESRI", auth_code = '54003', resampleAlg = "Bilinear", data_type = "Int16", srcNodata = -32768, dstNodata = -999) -> bool:
    dtt = gdal_datatypes()

    report(f"\rresampling {original_file}                                             ")

    ds = gdal.Warp(destination_file, original_file, dstSRS=f'{authority}:{auth_code}', resampleAlg=f"{resampleAlg}", srcNodata=srcNodata, dstNodata=dstNodata, outputType=dtt[data_type], xRes=resolution, yRes=resolution)
    ds = None

    return True

def create_polygon_geodataframe(lat_list, lon_list, auth = 'EPSG', code = 4326):
    geometry_ = Polygon(zip(lon_list, lat_list))
    polygon = geopandas.GeoDataFrame(index=[0], crs=f"{auth}:{code}", geometry=[geometry_])

    return polygon

    
def reproject_raster(input_raster, epsg, output_raster, method = "mode"):
    command = "gdalwarp -overwrite -t_srs EPSG:" + str(epsg) + " -r " + method + " -of GTiff " + input_raster + " " + output_raster
    os.system(command)
    print("\t> reprojected {0} to {1}".format(input_raster, epsg))

def set_nodata(input_tif, output_tif, nodata = -999):
    command = f"gdal_translate -of GTiff -a_nodata {nodata} {input_tif} {output_tif}"
    os.system(command)
    print(f"\t> set no data from {input_tif} to {output_tif}")

def get_extents(raster):
    src = gdal.Open(raster)
    upper_lef_x, xres, xskew, upper_left_y, yskew, yres  = src.GetGeoTransform()
    lower_right_x = upper_lef_x + (src.RasterXSize * xres)
    lower_right_y = upper_left_y + (src.RasterYSize * yres)
    return upper_lef_x, lower_right_y, lower_right_x, upper_left_y

def file_name(path_, extension=True):
    if extension:
        fn = os.path.basename(path_)
    else:
        fn = os.path.basename(path_).split(".")[0]
    return(fn)

def read_from(filename, decode_codec = None, v=False):
    '''
    a function to read ascii files
    '''
    try:
        if not decode_codec is None: g = open(filename, 'rb')
        else: g = open(filename, 'r')
    except:
        print(
            "\t! error reading {0}, make sure the file exists".format(filename))
        return

    file_text = g.readlines()
    
    if not decode_codec is None: file_text = [line.decode(decode_codec) for line in file_text]

    if v:
        print("\t> read {0}".format(file_name(filename)))
    g.close
    return file_text


def error(text_):
    print("\t! {string_}".format(string_=text_))

def strip_characters(input_:str, characters:str):
    return ''.join([i for i in input_ if not i in characters]) 

def list_files(folder, extension="*"):
    if folder.endswith("/"):
        if extension == "*":
            list_of_files = glob(folder + "*")
        else:
            list_of_files = glob(folder + "*." + extension if not extension.startswith(".") else f".{extension}")
    else:
        if extension == "*":
            list_of_files = glob(folder + "/*")
        else:
            list_of_files = glob(folder + "/*." + extension if not extension.startswith(".") else f".{extension}")
    return list_of_files


def list_all_files(folder, extension="*"):
    list_of_files = []
    # Getting the current work directory (cwd)
    thisdir = folder

    # r=root, d=directories, f = files
    for r, d, f in os.walk(thisdir):
        for file in f:
            if extension == "*":
                list_of_files.append(os.path.join(r, file))
            elif "." in extension:
                if file.endswith(extension[1:]):
                    list_of_files.append(os.path.join(r, file))
                    # print(os.path.join(r, file))
            else:
                if file.endswith(extension):
                    list_of_files.append(os.path.join(r, file))
                    # print(os.path.join(r, file))

    return list_of_files


def list_folders(directory):
    """
    directory: string or pathlike object
    """
    all_dirs = os.listdir(directory)
    dirs = [dir_ for dir_ in all_dirs if os.path.isdir(
        os.path.join(directory, dir_))]
    return dirs


def plot(x_list, y_list, plot_type='line', x_label="", y_label="", title_=""):
    '''
    plot_type: line
    '''
    plt.plot(x_list, y_list)
    plt.title(title_)
    plt.xlabel(x_label)
    plt.ylabel(y_label)
    plt.show()


def show(text_, error=False, same_line=True):
    if error:
        if same_line:
            sys.stdout.write("\r! {string_}                                                                                        ".format(string_=text_))
            sys.stdout.flush()
        else:
            sys.stdout.write("\r! {string_}".format(string_=text_))
            sys.stdout.flush()
    else:
        if same_line:
            sys.stdout.write("\r> {string_}                                                                                         ".format(string_=text_))
            sys.stdout.flush()
        else:
            sys.stdout.write("\r> {string_}".format(string_=text_))
            sys.stdout.flush()



def copy_file(filename, destination_path, delete_source=False, v = False, replace = True):
    '''
    a function to copy files
    '''
    if not replace:
        if exists(destination_path):
            if v:
                print(f"\t - file exists, skipping")
            return

    if not exists(filename):
        if not v:
            return
        print("\t> The file you want to copy does not exist")
        print(f"\t    - {filename}\n")
        ans = input("\t> Press  E then ENTER to Exit or C then ENTER to continue: ")

        counter = 0
        while (not ans.lower() == "c") and (not ans.lower() == "e"):
            ans = input("\t> Please, press E then ENTER to Exit or C then ENTER to continue: ")
            if counter > 2:
                print("\t! Learn to read instrunctions!!!!")
            counter += 1

        if ans.lower() == 'e': quit()
        if ans.lower() == 'c':
            write_to("log.txt", f"{filename}\n", mode='append')
            return


    if v:
        if delete_source:
            print(f"\t - [{get_file_size(filename)}] moving {filename} to \n\t\t{destination_path}")
        else:
            # print(f"\t - [{get_file_size(filename)}] copying {filename} to \n\t\t{destination_path}")

            sys.stdout.write('\rcopying ' + filename + '                        ')
            sys.stdout.flush()


    if not os.path.isdir(os.path.dirname(destination_path)):
        try:
            os.makedirs(os.path.dirname(destination_path))
        except:
            pass

    copyfile(filename, destination_path)
    if delete_source:
        try:
            os.remove(filename)
        except:
            error('coule not remove {fl}, make sure it is not in use'.format(fl=filename))

def copy_directory_tree(src, dst):
    if not os.path.exists(dst):
        os.makedirs(dst)
    for item in os.listdir(src):
        s = os.path.join(src, item)
        d = os.path.join(dst, item)
        if os.path.isdir(s):
            shutil.copytree(s, d, dirs_exist_ok=True)
        else:
            shutil.copy2(s, d)

def remove_header_duplicates(file_path, line_index = 0, separator = None):
    header = read_from(file_path)[line_index]
    if not separator is None:
        header = header.split(separator)
    else:
        header = header.split()
    # Rename duplicate columns
    column_counts = defaultdict(int)
    modified_header = []
    for col_name in header:
        column_counts[col_name] += 1
        if column_counts[col_name] > 1:
            modified_header.append(f"{col_name}_{column_counts[col_name]}")
        else:
            modified_header.append(col_name)
    return modified_header



def get_swat_timeseries(file_path, header_index = 1, col_name = "flo_out", object_number = None, object_name = None, separator = None, skip_rows = 3):

    if not exists(file_path):
        print('! swat+ result file does not exist')
        return None
    
    modified_header = remove_header_duplicates(file_path, header_index, separator) + ['extra1', 'extra2']


    try:
        df              = pandas.read_csv(file_path, delim_whitespace = True, skiprows = skip_rows, names = modified_header, index_col=False)
    except:
        sys.stdout.write(f'\r! could not read {file_path} using pandas, check the number of columns              \n')
        sys.stdout.flush()
        return None

    df = df.drop(columns='extra1')
    df = df.drop(columns='extra2')

    # print(file_path)
    # print(df)
    df['date']  = pandas.to_datetime(dict(year=df.yr, month=df.mon, day=df.day))

    if not object_number is None:
        df          = df[df['unit']== object_number]
    
    if not object_name is None:
        df          = df[df['name']== object_name]
    
    if (not col_name is None) and (not col_name == "*"):
        df      = df[['date', col_name]]

    return df


def clip_raster(dstDS, srcDS, cutline):
    if not os.path.isdir(os.path.dirname(dstDS)): os.makedirs(os.path.dirname(dstDS))
    options = gdal.WarpOptions(cutlineLayer = f'{cutline}', multithread = True, cropToCutline = True)
    gdal.Warp(dstDS, srcDS, options = options)
    return True

# def clip_raster(input_tif, clip_shapefile, output_tif, target_resolution=None, nodata=-999, tmp_fn = "tmp.dll", data_type="Int32", resampling="average", report=False, stealth=True):
#     """
#     nodata as string
#     """
#     import random
#     import string
#     tmp_fn = ''.join(random.choice(string.ascii_uppercase + string.ascii_lowercase + string.digits) for _ in range(16))
#     tmp_fn = f"tmp/{tmp_fn}.dll"
#     create_path("./tmp/")
#     if platform.system() == "Windows":
#         ctypes.windll.kernel32.SetFileAttributesW("tmp", 2)

#     if not os.path.isdir(os.path.dirname(output_tif)):
#         os.makedirs(os.path.dirname(output_tif))

#     if stealth:
#         if target_resolution is None:
#             os.system("gdalwarp -r {resample_type} -dstnodata {ds_nodata} -ot {dt_type} -of GTiff -cutline {shp_file} -crop_to_cutline {in_tif} {out_tif} -overwrite >> {tmp}".format(
#                 resample_type=resampling,
#                 ds_nodata=nodata,
#                 dt_type=data_type,
#                 shp_file=clip_shapefile,
#                 in_tif=input_tif,
#                 out_tif=output_tif,
#                 tmp = tmp_fn,
#             ))
#         else:
#             os.system("gdalwarp -r {resample_type} -tr {res_tar} -{res_tar} -dstnodata {ds_nodata} -ot {dt_type} -of GTiff -cutline {shp_file} -crop_to_cutline {in_tif} {out_tif} -overwrite >> {tmp}".format(
#                 resample_type=resampling,
#                 res_tar=target_resolution,
#                 ds_nodata=nodata,
#                 dt_type=data_type,
#                 shp_file=clip_shapefile,
#                 in_tif=input_tif,
#                 out_tif=output_tif,
#                 tmp = tmp_fn,
#             ))
#         try:
#             os.remove(tmp_fn)
#         except:
#             pass
#     else:
#         if target_resolution is None:
#             os.system("gdalwarp -r {resample_type} -dstnodata {ds_nodata} -ot {dt_type} -of GTiff -cutline {shp_file} -crop_to_cutline {in_tif} {out_tif} -overwrite".format(
#                 resample_type=resampling,
#                 ds_nodata=nodata,
#                 dt_type=data_type,
#                 shp_file=clip_shapefile,
#                 in_tif=input_tif,
#                 out_tif=output_tif,
#             ))
#         else:
#             os.system("gdalwarp -r {resample_type} -tr {res_tar} -{res_tar} -dstnodata {ds_nodata} -ot {dt_type} -of GTiff -cutline {shp_file} -crop_to_cutline {in_tif} {out_tif} -overwrite".format(
#                 resample_type=resampling,
#                 res_tar=target_resolution,
#                 ds_nodata=nodata,
#                 dt_type=data_type,
#                 shp_file=clip_shapefile,
#                 in_tif=input_tif,
#                 out_tif=output_tif,
#             ))

#     if report:
#         print("\t - clipped " + input_tif + " with " +
#               clip_shapefile + " and saved to\n\t  " + output_tif)

this_dir = os.getcwd()

def copy_projection(input_proj_tif, out_tif):
    dataset = gdal.Open(input_proj_tif)
    projection = dataset.GetProjection()
    geotransform = dataset.GetGeoTransform()
    if projection is None and geotransform is None:
        print('No projection or geotransform found on file' + input)
        sys.exit(1)
    dataset2 = gdal.Open(out_tif, gdal.GA_Update)
    if dataset2 is None:
        print('Unable to open', out_tif, 'for writing')
        sys.exit(1)
    if geotransform is not None:
        dataset2.SetGeoTransform(geotransform)
    if projection is not None:
        dataset2.SetProjection(projection)


def set_tif_nodata(in_tif, out_tif, no_data=-999, stealth = True):
    create_path(os.path.dirname(out_tif))
    create_path("tmp")
    hide_folder("tmp")
    if stealth:
        os.system("gdal_translate -of GTiff -a_nodata {2} {0} {1} > tmp/tmp.txt".format(in_tif, out_tif, no_data))
    else:
        os.system("gdal_translate -of GTiff -a_nodata {2} {0} {1}".format(in_tif, out_tif, no_data))


def clip_features(mask, input_feature, output_feature, keep_only_types = None, v = False) -> geopandas.GeoDataFrame:
    '''
    keep_only_types = ['MultiPolygon', 'Polygon', 'Point', etc]
    
    '''
    mask_gdf = geopandas.read_file(mask)
    input_gdf = geopandas.read_file(input_feature)
    create_path(output_feature)
    out_gdf = input_gdf.clip(mask_gdf.to_crs(input_gdf.crs))

    if not keep_only_types is None:
        out_gdf = out_gdf[out_gdf.geometry.apply(lambda x : x.type in keep_only_types)]

    out_gdf.to_file(output_feature)

    if v:
        print("\t  - clipped feature to " + output_feature)
    return out_gdf


def assign_default_projection(in_raster, out_raster, min_lon=-180, max_lon=180, min_lat=-90, max_lat=90):
    os.system("gdal_translate -a_srs WGS84 -a_ullr {min_lon} {max_lat} {max_lon} {min_lat} {in_raster} {out_raster}".format(
        min_lon=min_lon,
        max_lon=max_lon,
        min_lat=min_lat,
        max_lat=max_lat,
        in_raster=in_raster,
        out_raster=out_raster,
    ))
